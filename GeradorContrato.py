import PySimpleGUI as sg
from docx import Document
from datetime import datetime
import json
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate
from reportlab.platypus import Paragraph
from reportlab.lib.styles import getSampleStyleSheet

# Função para salvar o Número de Registro
def salvar_num_registro(num):
    with open('config.json', 'w') as file:
        json.dump({'num': num}, file)

# Função para carregar o Número de Registro
def carregar_num_registro():
    try:
        with open('config.json', 'r') as file:
            dado = json.load(file)
            if 'num' in dado:
                return int(dado['num'])
    except FileNotFoundError:
        print(f"Arquivo 'config.json' não encontrado.")
        return None

def convert_docx_to_pdf(docx_file, pdf_file):
    doc = Document(docx_file)
    
    # Criar um arquivo PDF
    pdf = SimpleDocTemplate(pdf_file, pagesize=letter)
    
    # Estilos para o PDF
    styles = getSampleStyleSheet()
    style_normal = styles['Normal']
    
    # Conteúdo do PDF
    content = []
    for para in doc.paragraphs:
        content.append(Paragraph(para.text, style_normal))
    
    # Adicionar o conteúdo ao PDF
    pdf.build(content)

sg.theme('LightBlue')

layout = [
    [sg.Text('Número da Venda', expand_x=4), sg.Input(size=(30), justification='center', key='num_venda')],
    [sg.Text('Empreedimento', expand_x=4), sg.Input(size=(30), justification='center', key='empreed')],
    [sg.Text('Quadra', expand_x=4), sg.Input(size=(30), justification='center', key='quadra')],
    [sg.Text('Lote', expand_x=4), sg.Input(size=(30), justification='center', key='lote')],
    [sg.Text('Cliente', expand_x=4), sg.Input(size=(30), justification='center', key='cliente')],
    [sg.Text('Observação(s)', expand_x=4)],
    [sg.Input(size=(30), key='obs')],
    [sg.Button('Limpar')],
    [sg.HSep(color='black', pad=(10, 10))],
    [sg.Button('Gerar Contrato', key='gerar')],
    [sg.Button('Gerar PDF', key='pdf')]
]

janela = sg.Window("Gerador de Contratos", layout, size=(400,300), element_justification='center')
 
while True:
    event, values = janela.read()

    if event == sg.WIN_CLOSED:
        break
    
    doc = Document("contratos.docx")

    if event == 'gerar':
        ## Pegando data
        data = datetime.now()
        dia = data.day
        mes = data.month
        ano = data.year

        n_registro = carregar_num_registro()
        
        num_venda = values['num_venda']
        empreed = values['empreed']
        quadra = values['quadra']
        lote = values['lote']
        cliente = values['cliente']
        obs = values['obs']

        doc.add_paragraph(f"Número da Venda: {num_venda} - Número de Registro: {n_registro} - {dia}/{mes}/{ano}")
        doc.add_paragraph(f"Empreendimento: {empreed} - Quadra: {quadra} - Lote: {lote}")
        doc.add_paragraph(f"Cliente: {cliente}")
        doc.add_paragraph(f"Obs: {obs}")
        doc.add_paragraph('------------------------------------------')
        sg.popup(f'Contrato nº {n_registro} gerado com sucesso!')
        salvar_num_registro(n_registro + 1)
        janela['num_venda'].update('')
        janela['empreed'].update('')
        janela['quadra'].update('')
        janela['lote'].update('')
        janela['cliente'].update('')
        janela['obs'].update('')

    doc.save("contratos.docx")

    if event == 'Limpar':
        janela['num_venda'].update('')
        janela['empreed'].update('')
        janela['quadra'].update('')
        janela['lote'].update('')
        janela['cliente'].update('')
        janela['obs'].update('')

    if event == 'pdf':
        convert_docx_to_pdf('contratos.docx', f'Contratos - {dia}.{mes}.{ano}.pdf')

janela.close()